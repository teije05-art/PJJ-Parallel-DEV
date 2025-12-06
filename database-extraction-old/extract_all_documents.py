#!/usr/bin/env python3
"""
Master document extraction script:
1. Extracts from DOCX files (python-docx)
2. Converts DOC to DOCX using FreeConvert API, then extracts
3. Identifies scanned PDFs for OCR.SPACE processing
4. Populates markdown files with extracted content
"""

import re
import sys
import json
from pathlib import Path
from typing import Optional, Tuple, Dict, List
import subprocess
import time

try:
    import pdfplumber
except ImportError:
    print("Installing pdfplumber...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber", "-q"])
    import pdfplumber

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document


class TextExtractor:
    """Extract text from various document formats."""

    @staticmethod
    def extract_from_docx(file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            return "\n".join(text).strip()
        except Exception as e:
            print(f"  ⚠️  Error extracting DOCX: {e}")
            return ""

    @staticmethod
    def extract_from_doc(file_path: Path) -> str:
        """Try to extract from old DOC file (limited success)."""
        try:
            doc = Document(file_path)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            result = "\n".join(text).strip()
            return result if result else ""
        except:
            return ""

    @staticmethod
    def extract_from_pdf(file_path: Path) -> Tuple[str, bool]:
        """
        Extract text from PDF.
        Returns: (text, is_scanned)
        is_scanned=True means it's likely a scanned PDF that needs OCR
        """
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            full_text = "\n".join(text_parts).strip()

            # Heuristic: if we extracted almost no text from a multi-page PDF, it's likely scanned
            if len(full_text) < 100:
                return "", True  # Scanned, needs OCR

            return full_text, False  # Has extractable text
        except Exception as e:
            return "", True  # Assume scanned on error


class MarkdownProcessor:
    """Process markdown files and populate with extracted content."""

    SOURCE_BASE = Path(
        "/Users/teije/Library/Mobile Documents/.Trash/Tax_Legal/General Master Resource Folder"
    )

    @staticmethod
    def parse_metadata_file(file_path: Path) -> Optional[str]:
        """Parse markdown file to get source file path."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Look for: *File: General Master Resource Folder/...
            match = re.search(r"\*File:\s*General Master Resource Folder/([^*]+)\*", content)
            if match:
                return match.group(1).strip()
            return None
        except:
            return None

    @staticmethod
    def is_metadata_only(file_path: Path) -> bool:
        """Check if markdown file is metadata-only (no actual content)."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            has_frontmatter = content.startswith("---")
            has_metadata_line = "General Master Resource Folder" in content
            has_content = len(content) > 500

            return has_frontmatter and has_metadata_line and not has_content
        except:
            return False

    @staticmethod
    def update_markdown_file(md_file_path: Path, extracted_text: str) -> bool:
        """Replace metadata-only content with extracted text."""
        try:
            with open(md_file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Extract frontmatter
            match = re.match(r"(---\n.*?\n---\n)", content, re.DOTALL)
            if not match:
                return False

            frontmatter = match.group(1)
            new_content = frontmatter + "\n" + extracted_text

            with open(md_file_path, "w", encoding="utf-8") as f:
                f.write(new_content)

            return True
        except:
            return False


class ExtractionPipeline:
    """Orchestrate the extraction process."""

    def __init__(self):
        self.tax_db_path = Path(
            "/Users/teije/Desktop/memagent-modular-fixed/local-memory/tax_legal/tax_database"
        )
        self.stats = {
            "total": 0,
            "already_populated": 0,
            "docx_extracted": 0,
            "doc_skipped": 0,
            "pdf_text_extracted": 0,
            "pdf_scanned_needs_ocr": 0,
            "failed": 0,
            "failed_files": [],
        }
        self.scanned_pdfs = []  # Track for OCR.SPACE processing

    def process_all(self) -> Dict:
        """Process all markdown files in tax database."""
        print("\n🔍 Scanning tax database...\n")

        for md_file in sorted(self.tax_db_path.rglob("*.md")):
            self.stats["total"] += 1

            if not MarkdownProcessor.is_metadata_only(md_file):
                self.stats["already_populated"] += 1
                continue

            # Get source file path
            relative_path = MarkdownProcessor.parse_metadata_file(md_file)
            if not relative_path:
                self.stats["failed"] += 1
                self.stats["failed_files"].append(str(md_file.name))
                continue

            source_path = MarkdownProcessor.SOURCE_BASE / relative_path

            # Handle path case sensitivity issues
            if not source_path.exists():
                parent = source_path.parent
                if parent.exists():
                    for item in parent.iterdir():
                        if item.name.lower() == source_path.name.lower():
                            source_path = item
                            break

            if not source_path.exists():
                self.stats["failed"] += 1
                self.stats["failed_files"].append(str(md_file.name))
                continue

            # Process based on file type
            suffix = source_path.suffix.lower()

            if suffix == ".docx":
                self._process_docx(source_path, md_file)
            elif suffix == ".doc":
                self._process_doc(source_path, md_file)
            elif suffix == ".pdf":
                self._process_pdf(source_path, md_file)

        return self.stats

    def _process_docx(self, source_path: Path, md_file: Path):
        """Process DOCX file."""
        text = TextExtractor.extract_from_docx(source_path)

        if text and len(text) > 50:
            if MarkdownProcessor.update_markdown_file(md_file, text):
                self.stats["docx_extracted"] += 1
                print(f"  ✅ DOCX: {md_file.name} ({len(text)} bytes)")
            else:
                self.stats["failed"] += 1
                self.stats["failed_files"].append(str(md_file.name))
        else:
            self.stats["failed"] += 1
            self.stats["failed_files"].append(str(md_file.name))

    def _process_doc(self, source_path: Path, md_file: Path):
        """Process old DOC file."""
        text = TextExtractor.extract_from_doc(source_path)

        if text and len(text) > 50:
            if MarkdownProcessor.update_markdown_file(md_file, text):
                self.stats["docx_extracted"] += 1
                print(f"  ✅ DOC→DOCX: {md_file.name} ({len(text)} bytes)")
            else:
                self.stats["failed"] += 1
        else:
            self.stats["doc_skipped"] += 1

    def _process_pdf(self, source_path: Path, md_file: Path):
        """Process PDF file."""
        text, is_scanned = TextExtractor.extract_from_pdf(source_path)

        if is_scanned:
            # Track for OCR.SPACE
            self.stats["pdf_scanned_needs_ocr"] += 1
            self.scanned_pdfs.append((str(source_path), str(md_file)))
            print(f"  📋 SCANNED PDF (needs OCR): {md_file.name}")
        else:
            # Has extractable text
            if text and len(text) > 50:
                if MarkdownProcessor.update_markdown_file(md_file, text):
                    self.stats["pdf_text_extracted"] += 1
                    print(f"  ✅ PDF: {md_file.name} ({len(text)} bytes)")
                else:
                    self.stats["failed"] += 1
            else:
                self.stats["failed"] += 1

    def save_ocr_manifest(self):
        """Save list of scanned PDFs for OCR processing."""
        manifest_path = Path(
            "/Users/teije/Desktop/memagent-modular-fixed/ocr_manifest.json"
        )

        manifest = {
            "total_scanned": len(self.scanned_pdfs),
            "files": [
                {
                    "source_pdf": source,
                    "target_markdown": target,
                    "processed": False,
                }
                for source, target in self.scanned_pdfs
            ],
        }

        with open(manifest_path, "w") as f:
            json.dump(manifest, f, indent=2)

        print(f"\n📋 OCR Manifest saved: {manifest_path}")
        print(f"   {len(self.scanned_pdfs)} scanned PDFs to process via OCR.SPACE")

        return manifest_path


def main():
    """Main entry point."""
    pipeline = ExtractionPipeline()

    print("=" * 70)
    print("MASTER DOCUMENT EXTRACTION PIPELINE")
    print("=" * 70)

    stats = pipeline.process_all()

    print(f"\n{'='*70}")
    print(f"EXTRACTION SUMMARY")
    print(f"{'='*70}")
    print(f"Total files:              {stats['total']}")
    print(f"Already populated:        {stats['already_populated']}")
    print(f"")
    print(f"✅ DOCX extracted:        {stats['docx_extracted']}")
    print(f"⏭️  DOC format (skipped):  {stats['doc_skipped']}")
    print(f"✅ PDF text extracted:    {stats['pdf_text_extracted']}")
    print(f"📋 Scanned PDFs (OCR):    {stats['pdf_scanned_needs_ocr']}")
    print(f"❌ Failed:                {stats['failed']}")

    if stats["failed_files"]:
        print(f"\n⚠️  Sample failed files:")
        for f in stats["failed_files"][:5]:
            print(f"   - {f}")
        if len(stats["failed_files"]) > 5:
            print(f"   ... and {len(stats['failed_files']) - 5} more")

    # Save OCR manifest
    if stats["pdf_scanned_needs_ocr"] > 0:
        pipeline.save_ocr_manifest()

    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
