#!/usr/bin/env python3
"""
Extract text from PDFs and Word documents and populate markdown files.
This script processes all metadata-only markdown files in the tax database
and extracts content from their source documents.
"""

import re
import sys
from pathlib import Path
from typing import Optional, Tuple

try:
    import pdfplumber
except ImportError:
    print("Installing pdfplumber...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber", "-q"])
    import pdfplumber

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document


class TextExtractor:
    """Extract text from PDF and Word documents."""

    @staticmethod
    def extract_from_pdf(file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return "\n".join(text).strip()
        except Exception as e:
            print(f"  ⚠️  Error extracting PDF: {e}")
            return ""

    @staticmethod
    def extract_from_docx(file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            return "\n".join(text).strip()
        except Exception as e:
            print(f"  ⚠️  Error extracting DOCX: {e}")
            return ""

    @staticmethod
    def extract_from_doc(file_path: Path) -> str:
        """Extract text from old DOC file using python-docx."""
        try:
            doc = Document(file_path)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            result = "\n".join(text).strip()
            return result if result else ""
        except Exception as e:
            # python-docx couldn't read it - will skip
            return ""

    @staticmethod
    def extract_text(file_path: Path) -> str:
        """Extract text from PDF or Word document."""
        if not file_path.exists():
            print(f"  ❌ File not found: {file_path}")
            return ""

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return TextExtractor.extract_from_pdf(file_path)
        elif suffix == ".docx":
            return TextExtractor.extract_from_docx(file_path)
        elif suffix == ".doc":
            return TextExtractor.extract_from_doc(file_path)
        else:
            print(f"  ⚠️  Unsupported file type: {suffix}")
            return ""


class MarkdownProcessor:
    """Process markdown files and extract/populate content."""

    SOURCE_BASE = Path(
        "/Users/teije/Library/Mobile Documents/.Trash/Tax_Legal/General Master Resource Folder"
    )

    @staticmethod
    def parse_metadata_file(file_path: Path) -> Optional[Tuple[str, str]]:
        """Parse metadata-only markdown file to get source file path."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Look for the metadata line: *File: General Master Resource Folder/...
            match = re.search(r"\*File:\s*General Master Resource Folder/([^*]+)\*", content)
            if match:
                relative_path = match.group(1).strip()
                return relative_path
            return None
        except Exception as e:
            print(f"  ❌ Error reading {file_path}: {e}")
            return None

    @staticmethod
    def is_metadata_only(file_path: Path) -> bool:
        """Check if markdown file is metadata-only (no actual content)."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Check if file only has frontmatter + metadata line
            has_frontmatter = content.startswith("---")
            has_metadata_line = "General Master Resource Folder" in content
            has_content = len(content) > 500  # Rough heuristic

            return has_frontmatter and has_metadata_line and not has_content
        except:
            return False

    @staticmethod
    def update_markdown_file(
        md_file_path: Path, extracted_text: str, original_format: str
    ) -> bool:
        """Replace metadata-only content with extracted text."""
        try:
            with open(md_file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # Extract frontmatter
            match = re.match(r"(---\n.*?\n---\n)", content, re.DOTALL)
            if not match:
                print(f"  ❌ Could not find frontmatter in {md_file_path}")
                return False

            frontmatter = match.group(1)

            # Create new content: frontmatter + extracted text
            new_content = frontmatter + "\n" + extracted_text

            # Write back
            with open(md_file_path, "w", encoding="utf-8") as f:
                f.write(new_content)

            return True
        except Exception as e:
            print(f"  ❌ Error updating {md_file_path}: {e}")
            return False

    @staticmethod
    def process_directory(tax_db_path: Path) -> dict:
        """Process all markdown files in tax database."""
        stats = {
            "total_files": 0,
            "metadata_only": 0,
            "already_populated": 0,
            "extracted": 0,
            "failed": 0,
            "failed_files": [],
        }

        for md_file in sorted(tax_db_path.rglob("*.md")):
            stats["total_files"] += 1

            if not MarkdownProcessor.is_metadata_only(md_file):
                stats["already_populated"] += 1
                continue

            stats["metadata_only"] += 1

            # Parse metadata
            relative_path = MarkdownProcessor.parse_metadata_file(md_file)
            if not relative_path:
                stats["failed"] += 1
                stats["failed_files"].append(str(md_file))
                continue

            # Find source file
            source_path = MarkdownProcessor.SOURCE_BASE / relative_path

            # Try alternative path resolution (handle spaces and special chars)
            if not source_path.exists():
                # Try with exact match ignoring case
                parent = source_path.parent
                if parent.exists():
                    for item in parent.iterdir():
                        if item.name.lower() == source_path.name.lower():
                            source_path = item
                            break

            if not source_path.exists():
                print(f"  ⚠️  Skipping {md_file.name} - source not found: {source_path}")
                stats["failed"] += 1
                stats["failed_files"].append(str(md_file))
                continue

            # Extract text
            extracted_text = TextExtractor.extract_text(source_path)

            if not extracted_text:
                stats["failed"] += 1
                stats["failed_files"].append(str(md_file))
                continue

            # Update markdown
            if MarkdownProcessor.update_markdown_file(md_file, extracted_text, source_path.suffix):
                stats["extracted"] += 1
                print(
                    f"  ✅ {md_file.name} ({len(extracted_text)} bytes)"
                )
            else:
                stats["failed"] += 1
                stats["failed_files"].append(str(md_file))

        return stats


def main():
    """Main entry point."""
    tax_db_path = Path(
        "/Users/teije/Desktop/memagent-modular-fixed/local-memory/tax_legal/tax_database"
    )

    if not tax_db_path.exists():
        print(f"❌ Tax database not found: {tax_db_path}")
        sys.exit(1)

    print("\n🔍 Scanning tax database for metadata-only files...\n")

    stats = MarkdownProcessor.process_directory(tax_db_path)

    print(f"\n{'='*70}")
    print(f"EXTRACTION SUMMARY")
    print(f"{'='*70}")
    print(f"Total files:          {stats['total_files']}")
    print(f"Already populated:    {stats['already_populated']}")
    print(f"Metadata-only found:  {stats['metadata_only']}")
    print(f"Successfully extracted: {stats['extracted']}")
    print(f"Failed:               {stats['failed']}")

    if stats["failed_files"]:
        print(f"\n⚠️  Failed files:")
        for f in stats["failed_files"][:10]:  # Show first 10
            print(f"   - {f}")
        if len(stats["failed_files"]) > 10:
            print(f"   ... and {len(stats['failed_files']) - 10} more")

    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
